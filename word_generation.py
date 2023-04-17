import docx
import os
from docx.shared import Cm
from PIL import Image
from docx.enum.text import WD_ALIGN_PARAGRAPH


def word_generation(image, text, template, output, name):
    # Open the image and rotate it by 90 degrees
    image = Image.open(image)
    rotated_image = image.rotate(90, resample=Image.BICUBIC, expand=True)
    rotated_image.save(name + ".png")

    # Open the Word document
    doc = docx.Document(template)

    # Get the table in the document (assuming it's the first table)
    table = doc.tables[0]

    # Add an image
    cell_image = table.cell(3, 0)

    # Create a new paragraph in the cell
    cell_image._element.clear_content()
    paragraph = cell_image.add_paragraph()

    # Add a run with the image
    image_path = name + ".png"
    run = paragraph.add_run()
    run.add_picture(image_path, height=Cm(23))
    cell_image.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_image.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add some text
    cell_text = table.cell(2, 2)
    cell_text._element.clear_content()
    paragraph_text = cell_text.add_paragraph()
    paragraph_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_text.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
    run = paragraph_text.add_run(text)
    run.font.name = "Arial"
    run.font.size = docx.shared.Pt(11)

    # Save the modified document
    doc.save(output)

    # Delete the temporary file
    os.remove(name + ".png")
